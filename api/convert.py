"""Convert markdown to DOCX or PDF.

DOCX: pure python-docx, no external binaries.
PDF:  ReportLab, no external binaries or GTK/Pango DLLs required on Windows.
"""
from __future__ import annotations

from html import escape
import io
import re


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def to_docx(markdown: str) -> bytes:
    from docx import Document  # python-docx

    doc = Document()
    lines = markdown.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]

        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("|"):
            # Collect consecutive table lines, skip separator rows
            table_lines: list[str] = []
            while i < len(lines) and lines[i].startswith("|"):
                if not re.match(r"^\|[-| :]+\|$", lines[i]):
                    table_lines.append(lines[i])
                i += 1
            i -= 1  # outer loop will increment

            if table_lines:
                rows = [
                    [c.strip() for c in row.split("|")[1:-1]]
                    for row in table_lines
                ]
                n_cols = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=n_cols)
                tbl.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx in range(n_cols):
                        txt = row[c_idx] if c_idx < len(row) else ""
                        txt = re.sub(r"\*\*(.+?)\*\*", r"\1", txt)
                        txt = re.sub(r"\*(.+?)\*", r"\1", txt)
                        tbl.rows[r_idx].cells[c_idx].text = txt
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _runs(p, line[2:].strip())
        elif re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            _runs(p, re.sub(r"^\d+\.\s+", "", line))
        elif line.strip() in ("---", "***", "___"):
            doc.add_paragraph("─" * 60)
        elif line.startswith("> "):
            doc.add_paragraph(line[2:].strip())
        elif line.strip():
            p = doc.add_paragraph()
            _runs(p, line)

        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _runs(paragraph, text: str) -> None:
    """Parse **bold**, *italic*, `code` and add styled runs."""
    pattern = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            paragraph.add_run(text[last : m.start()])
        if m.group(1):
            paragraph.add_run(m.group(1)).bold = True
        elif m.group(2):
            paragraph.add_run(m.group(2)).italic = True
        elif m.group(3):
            r = paragraph.add_run(m.group(3))
            r.font.name = "Courier New"
        last = m.end()
    if last < len(text):
        paragraph.add_run(text[last:])


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def to_pdf(markdown: str) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "ReviewBody",
        parent=styles["BodyText"],
        fontName="Times-Roman",
        fontSize=10.5,
        leading=15,
        spaceAfter=7,
        alignment=TA_LEFT,
    )
    heading_styles = {
        1: ParagraphStyle("ReviewHeading1", parent=styles["Heading1"], fontSize=18, leading=22, spaceAfter=10),
        2: ParagraphStyle("ReviewHeading2", parent=styles["Heading2"], fontSize=14, leading=18, spaceBefore=8, spaceAfter=6),
        3: ParagraphStyle("ReviewHeading3", parent=styles["Heading3"], fontSize=12, leading=15, spaceBefore=6, spaceAfter=4),
        4: ParagraphStyle("ReviewHeading4", parent=styles["Heading4"], fontSize=10.5, leading=13, spaceBefore=4, spaceAfter=3),
    }
    code = ParagraphStyle(
        "ReviewCode",
        parent=body,
        fontName="Courier",
        fontSize=8.5,
        leading=11,
        leftIndent=12,
        backColor=colors.whitesmoke,
        borderPadding=5,
    )
    quote = ParagraphStyle(
        "ReviewQuote",
        parent=body,
        leftIndent=14,
        textColor=colors.HexColor("#555555"),
        borderColor=colors.lightgrey,
        borderWidth=1,
        borderPadding=5,
    )

    story = []
    lines = markdown.splitlines()
    i = 0
    in_code_block = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                story.append(Paragraph("<br/>".join(escape(part) for part in code_lines), code))
                story.append(Spacer(1, 6))
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped in ("---", "***", "___"):
            story.append(Spacer(1, 8))
            i += 1
            continue

        heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            story.append(Paragraph(_inline_to_reportlab(heading_match.group(2)), heading_styles[level]))
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = lines[i].strip()
                if not re.match(r"^\|[-| :]+\|$", row):
                    table_lines.append(row)
                i += 1
            if table_lines:
                story.append(_markdown_table_to_reportlab(table_lines, body))
                story.append(Spacer(1, 8))
            continue

        bullet_match = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet_match:
            items = []
            while i < len(lines):
                match = re.match(r"^[-*]\s+(.+)$", lines[i].strip())
                if not match:
                    break
                items.append(ListItem(Paragraph(_inline_to_reportlab(match.group(1)), body)))
                i += 1
            story.append(ListFlowable(items, bulletType="bullet", leftIndent=18))
            story.append(Spacer(1, 5))
            continue

        number_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if number_match:
            items = []
            while i < len(lines):
                match = re.match(r"^\d+\.\s+(.+)$", lines[i].strip())
                if not match:
                    break
                items.append(ListItem(Paragraph(_inline_to_reportlab(match.group(1)), body)))
                i += 1
            story.append(ListFlowable(items, bulletType="1", leftIndent=18))
            story.append(Spacer(1, 5))
            continue

        if stripped.startswith("> "):
            story.append(Paragraph(_inline_to_reportlab(stripped[2:]), quote))
            i += 1
            continue

        story.append(Paragraph(_inline_to_reportlab(stripped), body))
        i += 1

    if in_code_block and code_lines:
        story.append(Paragraph("<br/>".join(escape(part) for part in code_lines), code))

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        rightMargin=0.8 * inch,
        leftMargin=0.8 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
    )
    doc.build(story or [Paragraph("No review content was provided.", body)])
    return buf.getvalue()


def _inline_to_reportlab(text: str) -> str:
    escaped = escape(text)
    escaped = re.sub(r"`([^`]+)`", r'<font name="Courier">\1</font>', escaped)
    escaped = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", escaped)
    escaped = re.sub(r"\*([^*]+)\*", r"<i>\1</i>", escaped)
    return escaped


def _markdown_table_to_reportlab(table_lines: list[str], body_style) -> Table:
    from reportlab.lib import colors
    from reportlab.platypus import TableStyle
    from reportlab.platypus import Paragraph, Table

    rows = [
        [_inline_to_reportlab(cell.strip()) for cell in row.split("|")[1:-1]]
        for row in table_lines
    ]
    n_cols = max(len(row) for row in rows)
    data = []
    for row in rows:
        padded = row + [""] * (n_cols - len(row))
        data.append([Paragraph(cell, body_style) for cell in padded])

    table = Table(data, repeatRows=1, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.35, colors.lightgrey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f1f5f9")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    return table
